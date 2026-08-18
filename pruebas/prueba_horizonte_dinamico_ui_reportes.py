"""Pruebas de horizonte dinámico, intervalos 80/95, UI y reportes."""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

import numpy as np
import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
sys.path.append(str(ROOT))

#: Peldanos de la escalera de confianza metodologica retirada en P0-G. Se
#: comprueban por AUSENCIA: fijar el conjunto nuevo reconstruiria la escalera
#: dentro de la prueba (REQ 5, REQ 7).
ESCALERA_RETIRADA = {
    "alto", "medio", "bajo", "no recomendable",
    "alta confiabilidad relativa", "proyeccion con cautela", "proyección con cautela",
    "proyectable con cautela", "solo proyeccion de corto plazo",
    "solo proyección de corto plazo", "escenario de alta incertidumbre",
}


def _sin_escalera_de_confianza(factibilidad: dict) -> None:
    """La confianza publicada debe ser descriptiva, nunca un grado ordinal."""
    confianza = str(factibilidad.get("nivel_confianza_metodologica") or "").strip().lower()
    assert confianza not in ESCALERA_RETIRADA, confianza


def _periodos(n: int) -> list[str]:
    return [f"{2021 + i // 12}_{i % 12 + 1}" for i in range(n)]


def _serie_estable(n: int = 120) -> pd.DataFrame:
    valores = [100.0 + 0.25 * i + 0.02 * np.sin(i / 4.0) for i in range(n)]
    return pd.DataFrame({"Periodo": _periodos(n), "Indice": valores})


def _serie_erratica(n: int = 72) -> pd.DataFrame:
    rng = np.random.default_rng(202)
    valores = [max(1.0, 100.0 + (35.0 if i % 3 == 0 else -25.0) + rng.normal(0, 18)) for i in range(n)]
    return pd.DataFrame({"Periodo": _periodos(n), "Indice": valores})


#: Fixtures versionados del caso 7777777 (hallazgo H-13).
#:
#: Hasta julio de 2026 esta prueba leia reportes_generados/7777777.csv **si
#: existia** y, si no, usaba una serie codificada aqui mismo. Las dos series no
#: son equivalentes: la del CSV selecciona Drift y la codificada selecciona Holt
#: lineal, de modo que el resultado de la suite dependia de si una ejecucion
#: anterior habia dejado el archivo en el disco.
#:
#: Ahora las dos series son fixtures explicitos y cada escenario tiene su propia
#: prueba. Nada se lee de carpetas de salida.
FIXTURES = ROOT / "tests" / "fixtures" / "horizonte_dinamico"


def _serie_fixture(nombre: str) -> pd.DataFrame:
    """Carga una serie de prueba versionada. Falla si el fixture no esta."""
    ruta = FIXTURES / nombre
    if not ruta.is_file():
        raise FileNotFoundError(
            f"Falta el fixture {ruta}. Las series de prueba estan versionadas: "
            "no se generan ni se leen de carpetas de salida."
        )
    serie = pd.read_csv(ruta, dtype={"Periodo": str})
    return serie[["Periodo", "Indice"]].reset_index(drop=True)


def _serie_777_drift() -> pd.DataFrame:
    """Serie del caso 7777777 que selecciona Drift."""
    return _serie_fixture("serie_drift.csv")


def _serie_777_holt() -> pd.DataFrame:
    """Serie del caso 7777777 que selecciona Holt lineal."""
    return _serie_fixture("serie_holt.csv")


def _evaluacion_horizonte(
    horizonte: int,
    *,
    tecnico: bool = True,
    escenario: bool = True,
    rmse: float = 1.0,
    mae: float = 0.8,
    mape: float = 1.0,
    ic95: float = 0.10,
    estado: str = "Proyección extendida con cautela",
    decision: str | None = None,
) -> dict:
    permitido = bool(tecnico or escenario)
    no_recomendable = not permitido
    return {
        "horizonte": horizonte,
        "permitido": permitido,
        "permitido_para_proyeccion_tecnica": tecnico,
        "permitido_como_escenario": escenario,
        "no_recomendable": no_recomendable,
        "decision": decision or ("Permitido para proyección técnica" if tecnico else ("Permitido solo como escenario de alta incertidumbre" if escenario else "No recomendable")),
        "estado": estado if permitido else "No recomendable para el horizonte solicitado",
        "tipo_uso": "Proyección extendida con cautela" if horizonte <= 12 else "Escenario estadístico extendido",
        "modelo": {"nombre_visible": "Drift"},
        "backtesting": {
            "metricas": {
                "rmse": rmse,
                "mae": mae,
                "mape": mape,
                "smape": mape,
                "sesgo_medio": 0.1,
                "porcentaje_errores_extremos": 0.0,
                "estabilidad_error": 0.2,
            }
        },
        "evaluacion_intervalos": {
            "ancho_relativo_80_maximo": ic95 / 2.0,
            "ancho_relativo_95_maximo": ic95,
        },
        "mensaje_horizonte": "backtesting e intervalos aceptables" if permitido else "IC95 relativo excesivo para el horizonte.",
    }


def test_horizonte_18_se_permite_como_escenario() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    resultado = ejecutar_proyeccion(_serie_estable(120), 2032, 6, 2021)
    assert resultado["proyeccion_generada"] is True
    assert resultado["horizonte_solicitado"] == 18
    assert resultado["horizonte_permitido"] == 18
    # P0-G, 13-08-2026. Antes se exigia que el estado perteneciera al conjunto
    # {"Escenario estadistico extendido", "Escenario de alta incertidumbre",
    # "Proyeccion extendida"}: tres peldanos de la escalera de confianza, cuyos
    # literales no tenian fuente que autorizara la consecuencia aplicada. La
    # escalera se retiro y `estado` es hoy descriptivo. Fijar el conjunto NUEVO
    # la reconstruiria dentro de la prueba; lo que se comprueba es que no vuelva.
    factibilidad = resultado["factibilidad"]
    assert str(factibilidad.get("estado") or "").strip(), factibilidad
    _sin_escalera_de_confianza(factibilidad)
    assert resultado["estado_metodologico"] != "resultado_metodologicamente_sustentado"
    assert {"P0-C", "P0-E"} <= set(resultado["bloqueos_metodologicos"] or {})
    assert {"limite_inferior_95", "limite_superior_95"}.issubset(resultado["proyecciones"].columns)


def test_horizonte_25_estable_se_permite_si_hay_evidencia() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    resultado = ejecutar_proyeccion(_serie_estable(120), 2033, 1, 2021)
    assert resultado["horizonte_solicitado"] == 25
    assert resultado["proyeccion_generada"] is True
    assert resultado["horizonte_permitido"] >= 18
    assert resultado["horizonte_info"]["accion"] in {"permitir como escenario", "permitir con cautela", "permitir"}


def test_horizonte_25_erratico_no_se_fuerza() -> None:
    """El horizonte pedido nunca se concede por encima del maximo admisible.

    CIERRE 08-08-2026: una serie erratica ya no se bloquea por la amplitud de su
    intervalo -ese corte no tenia fuente-, de modo que la proyeccion puede
    generarse. Lo que esta prueba vigila sigue en pie y es lo importante: que
    h=25 no se conceda solo porque se pidio. Si se entrega, debe ser porque el
    maximo admisible lo cubre, y el horizonte entregado nunca puede superarlo.
    """
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    resultado = ejecutar_proyeccion(_serie_erratica(72), 2029, 1, 2021)
    assert resultado["horizonte_solicitado"] == 25
    if not resultado["proyeccion_generada"]:
        return
    info = resultado.get("analisis_horizontes_completo") or {}
    admisible = int(info.get("horizonte_maximo_admisible") or 0)
    permitido = int(resultado["horizonte_permitido"])
    assert permitido <= admisible, (permitido, admisible)
    assert permitido <= 25


def test_umbral_h12_31_es_alta_incertidumbre_no_bloqueo() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import _evaluar_intervalos_prediccion

    tabla = pd.DataFrame(
        {
            "indice_proyectado": [100.0],
            "limite_inferior_95": [84.5],
            "limite_superior_95": [115.5],
            "ancho_relativo_95": [0.31],
            "ancho_relativo_80": [0.18],
        }
    )
    evaluacion = _evaluar_intervalos_prediccion(tabla, horizonte=12)
    assert evaluacion["critico"] is False
    assert evaluacion["clasificacion"] == "cautela"


def test_umbral_h18_535_no_recomienda() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import _evaluar_intervalos_prediccion

    tabla_ok = pd.DataFrame({"indice_proyectado": [100.0], "ancho_relativo_95": [0.40], "ancho_relativo_80": [0.24]})
    tabla_limite = pd.DataFrame({"indice_proyectado": [100.0], "ancho_relativo_95": [0.535], "ancho_relativo_80": [0.32]})
    tabla_no = pd.DataFrame({"indice_proyectado": [100.0], "ancho_relativo_95": [0.56], "ancho_relativo_80": [0.32]})
    assert _evaluar_intervalos_prediccion(tabla_ok, horizonte=18)["critico"] is False
    assert _evaluar_intervalos_prediccion(tabla_limite, horizonte=18)["critico"] is True
    assert _evaluar_intervalos_prediccion(tabla_no, horizonte=18)["critico"] is True


def test_horizonte_maximo_no_es_18_si_h18_no_cumple() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import determinar_horizonte_maximo_estadistico

    # El servicio evalua horizontes mes a mes; la conciliacion exige consecutividad.
    evaluaciones = [
        _evaluacion_horizonte(h, rmse=0.8 + 0.19 * h, mape=0.6 + 0.12 * h, ic95=0.05 + 0.01 * h)
        for h in range(1, 13)
    ]
    evaluaciones.append(
        _evaluacion_horizonte(
            18,
            tecnico=False,
            escenario=False,
            rmse=5.8823,
            mae=5.1511,
            mape=3.8063,
            ic95=0.5350,
            decision="No recomendable",
        )
    )
    info = determinar_horizonte_maximo_estadistico(None, None, None, evaluaciones, None, horizonte_solicitado=12)
    assert info["horizonte_maximo_recomendado"] == 12
    # Ningun horizonte quedo clasificado como escenario puro en este caso.
    assert info["horizonte_maximo_permitido_como_escenario"] == 0
    assert info["horizontes_no_recomendables"] == [18]
    assert "h=18" in info["mensaje_ui"]
    assert info["tipo_uso_recomendado"] == "Proyección extendida con cautela"


def test_horizontes_se_evaluan_mes_a_mes_y_no_solo_atajos_ui() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import _horizontes_evaluacion

    horizontes = _horizontes_evaluacion(12, 61)
    assert horizontes[:6] == (1, 2, 3, 4, 5, 6)
    assert 12 in horizontes
    assert 18 in horizontes
    assert set(horizontes) != {1, 3, 6, 12, 18}


def test_mensaje_no_recomendables_no_imprime_lista_vacia() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import determinar_horizonte_maximo_estadistico

    evaluaciones = [_evaluacion_horizonte(h, estado="Alta confiabilidad relativa") for h in range(1, 7)]
    info = determinar_horizonte_maximo_estadistico(None, None, None, evaluaciones, None, horizonte_solicitado=6)
    assert info["horizontes_no_recomendables"] == []
    assert "[]" not in info["mensaje_ui"]
    assert "No se identificaron horizontes no recomendables" in info["mensaje_no_recomendables"]
    assert info["primer_horizonte_no_viable"] == 0


def test_h18_no_se_redacta_como_intervalos_aceptables() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import determinar_horizonte_maximo_estadistico
    from app_icociv.reportes.generador_reportes import _lineas_horizontes, _lineas_determinacion_horizonte

    # El servicio evalua horizontes mes a mes; la conciliacion exige consecutividad.
    evaluaciones = [_evaluacion_horizonte(h) for h in range(1, 13)]
    evaluaciones.append(
        _evaluacion_horizonte(18, tecnico=False, escenario=False, rmse=5.8823, mape=3.8063, ic95=0.5350)
    )
    info = determinar_horizonte_maximo_estadistico(None, None, None, evaluaciones, None, horizonte_solicitado=12)
    resultado = {"horizonte_info": info, "horizonte_solicitado": 12, "model_name": "Drift", "horizonte_permitido": 12}
    lineas = "\n".join(_lineas_horizontes(resultado) + _lineas_determinacion_horizonte(resultado))
    assert "Horizonte máximo recomendado: 12" in lineas
    assert "h=18 no recomendable" in lineas
    assert "18 meses" in lineas
    linea_h18 = next(linea for linea in lineas.splitlines() if linea.startswith("18 meses |"))
    assert "permitido=No" in linea_h18
    assert "técnico=No" in linea_h18
    assert "escenario=No" in linea_h18
    assert "backtesting e intervalos aceptables" not in linea_h18


def test_el_modelo_entregado_es_el_de_minimo_error_oos_global() -> None:
    """CIERRE H-9, 08-08-2026: la seleccion deja de ponderar por 1/h.

    Esta prueba exigia `drift` por su nombre. Drift ganaba porque la regla
    anterior daba mas peso a los horizontes cortos, donde domina. Retirada esa
    ponderacion -sin fuente-, gana el de menor error cuadratico medio fuera de
    muestra global. En esta serie sintetica es `exponencial_log_lineal`, con
    RMSE global 1,44 frente a 1,95; drift sigue siendo tres veces mejor en h=1.

    Se fija la propiedad de la regla, no la identidad del ganador.
    """
    from app_icociv.proyeccion.servicio_proyeccion import (
        ejecutar_proyeccion,
        seleccionar_modelo_por_rmse_oos_global,
    )
    from app_icociv.validacion.backtesting import ejecutar_backtesting_comparativo

    valores = [100.0 + 0.45 * i + 1.8 * np.sin(i / 3.0) for i in range(61)]
    serie = pd.DataFrame({"Periodo": _periodos(61), "Indice": valores})
    resultado = ejecutar_proyeccion(serie, 2026, 4, 2021)
    horizontes = tuple(range(1, 25))
    esperado = seleccionar_modelo_por_rmse_oos_global(
        ejecutar_backtesting_comparativo(serie, horizontes=horizontes, anio_base=2021),
        horizontes,
    )
    assert esperado is not None
    assert resultado["modelo_codigo"] == esperado, (resultado["modelo_codigo"], esperado)


def test_combo_no_cambia_con_rueda_y_periodo_iso() -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    from PySide6.QtWidgets import QApplication
    from app_icociv.interfaz.ventana_principal import ComboBoxSinRueda, periodo_a_iso

    app = QApplication.instance() or QApplication([])
    combo = ComboBoxSinRueda()
    combo.addItems(["A", "B", "C"])
    combo.setCurrentIndex(1)

    class Evento:
        def __init__(self) -> None:
            self.ignorado = False

        def ignore(self) -> None:
            self.ignorado = True

    evento = Evento()
    combo.wheelEvent(evento)
    assert evento.ignorado is True
    assert combo.currentIndex() == 1
    assert periodo_a_iso("2025_1") == "2025-01"
    assert app is not None


def test_grafica_usa_periodos_reales() -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    from PySide6.QtWidgets import QApplication
    from app_icociv.interfaz.ventana_principal import VentanaPrincipal
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    from matplotlib.figure import Figure

    app = QApplication.instance() or QApplication([])
    serie = _serie_estable(36)
    resultado = ejecutar_proyeccion(serie, 2024, 2, 2021)
    ventana = VentanaPrincipal()
    etiquetas = ventana._dibujar_grafica(Figure(), serie, resultado)
    assert etiquetas[0] == "2021-01"
    assert all("-" in etiqueta for etiqueta in etiquetas[:5])
    assert app is not None


def test_contenedor_ruta_tiene_scroll_y_tooltip_hover_combo() -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    from PySide6.QtWidgets import QApplication, QScrollArea
    from app_icociv.interfaz.ventana_principal import VentanaPrincipal

    app = QApplication.instance() or QApplication([])
    ventana = VentanaPrincipal()
    assert isinstance(ventana.scroll_ruta_jerarquica, QScrollArea)
    assert ventana.scroll_ruta_jerarquica.maximumHeight() <= 150
    assert ventana.combo_grupo.view().hasMouseTracking() is True
    assert "scroll" in ventana.scroll_ruta_jerarquica.toolTip().lower()
    assert app is not None


def test_el_catalogo_se_define_por_estimabilidad_y_no_por_longitud() -> None:
    """AUDITORIA 09-08-2026 (P0-B). Antes fijaba el escalonado por «niveles».

    Ese escalonado lo gobernaban siete literales sin fuente -`horizonte>=7`,
    `volatilidad>0.035`, `n_obs>=48`, `24`, `volatilidad>0.05`, `horizonte<=6`-
    que decidian que modelos competian. Se retiraron.

    Se conserva lo que la prueba verificaba de fondo -que el catalogo no incluye
    modelos fuera de alcance- y se sustituye la expectativa de «niveles» por la
    propiedad vigente: el catalogo NO depende de la longitud de la serie ni del
    horizonte, y lo unico que excluye a un modelo ademas de no ser estimable es
    tener un parametro propio sin fuente.
    """
    from app_icociv.proyeccion.servicio_proyeccion import (
        MODELOS_PARAMETRO_SIN_SUSTENTO,
        _modelos_para_analisis,
    )

    serie_corta = _serie_estable(36)
    modelos_cortos, politica_corta = _modelos_para_analisis(
        serie_corta,
        horizonte_solicitado=3,
        validacion_serie={"observaciones": len(serie_corta)},
        outliers=[],
    )
    serie_larga = _serie_estable(96)
    modelos_largos, politica_larga = _modelos_para_analisis(
        serie_larga,
        horizonte_solicitado=18,
        validacion_serie={"observaciones": len(serie_larga)},
        outliers=[],
    )

    for modelos in (modelos_cortos, modelos_largos):
        assert "arima_simple" not in modelos
        assert "polinomico_2" not in modelos
        for excluido in MODELOS_PARAMETRO_SIN_SUSTENTO:
            assert excluido not in modelos, excluido

    # El catalogo ya no depende de la longitud de la serie ni del horizonte.
    assert modelos_cortos == modelos_largos, (modelos_cortos, modelos_largos)
    for politica in (politica_corta, politica_larga):
        assert "estimab" in politica["criterio_elegibilidad"].lower()
        assert "niveles_activados" not in politica


def test_no_se_aplican_ensambles_de_modelos() -> None:
    """El ensamble fue retirado: siempre se aplica un único modelo por serie."""
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    for serie, anio, mes in ((_serie_estable(72), 2027, 1), (_serie_estable(120), 2032, 6)):
        resultado = ejecutar_proyeccion(serie, anio, mes, 2021)
        assert resultado.get("modelo_codigo") != "ensamble_obras_civiles"
        assert "componentes_ensamble" not in resultado
        assert not resultado.get("es_ensamble")


def test_mase_auxiliar_y_primer_periodo_proyectado() -> None:
    from app_icociv.estadistica.metricas import calcular_mase
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    mase = calcular_mase([2, 4], [1, 5], [1, 2, 4])
    assert abs(mase - (1.0 / 1.5)) < 1e-9

    serie = pd.DataFrame({"Periodo": _periodos(61), "Indice": [100.0 + 0.2 * i for i in range(61)]})
    resultado = ejecutar_proyeccion(serie, 2026, 2, 2021)
    assert resultado["proyecciones"].iloc[0]["periodo"] == "2026_2"
    assert resultado["proyecciones"].iloc[0]["periodo"] != serie.iloc[-1]["Periodo"]
    if resultado["modelo_codigo"] != "naive":
        assert float(resultado["proyecciones"].iloc[0]["indice_proyectado"]) != float(serie.iloc[-1]["Indice"])


def test_graficas_reporte_y_fundamento_estadistico() -> None:
    from app_icociv.reportes import generador_reportes, graficas
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    resultado = ejecutar_proyeccion(_serie_estable(72), 2027, 6, 2021)

    # Tras el rediseño de informes, el backtesting se ilustra con el error por
    # horizonte en lugar de la serie observado-vs-predicho.
    serie = _serie_estable(72)
    assert graficas.grafica_principal(serie, resultado) is not None
    assert graficas.grafica_errores_horizonte(resultado) is not None

    resultado = ejecutar_proyeccion(_serie_estable(72), 2027, 6, 2021)
    fundamento = generador_reportes._lineas_fundamento_estadistico(resultado)
    referencias = generador_reportes._referencias_estadisticas(resultado)
    assert any("walk-forward" in linea for linea in fundamento)
    assert any("Hyndman" in ref for ref in referencias)
    assert any("DANE" in ref for ref in referencias)


def test_modelos_tradicionales_y_no_ejecutados_quedan_trazados() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    resultado = ejecutar_proyeccion(_serie_estable(96), 2030, 6, 2021)
    catalogo = {item["codigo"]: item for item in resultado.get("catalogo_modelos", [])}
    assert catalogo["lineal"]["ejecutado"] == "Si"
    assert "logaritmico" in catalogo
    assert "exponencial_log_lineal" in catalogo
    assert catalogo["exponencial_log_lineal"]["ejecutado"] in {"Si", "No"}
    assert catalogo["exponencial_log_lineal"]["razon"]
    assert "nivel" in catalogo["lineal"]["tipo"].lower()
    assert "log-variación" in catalogo["log_variacion"]["tipo"].lower()


def test_parametros_drift_receta_y_csv_reproducible() -> None:
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion
    from app_icociv.reportes.generador_reportes import (
        _lineas_parametros_reproducibles,
        _lineas_receta_reproduccion,
        construir_dataframe_reproducibilidad,
        generar_csv_reproducibilidad,
    )

    # CIERRE H-9, 08-08-2026: esta serie ya no selecciona `drift` sino `lineal`.
    # Antes lo hacia porque la seleccion ponderaba los horizontes por 1/h, donde
    # drift domina; ahora gana el de menor error cuadratico medio OOS global.
    # Lo que la prueba vigila -que los parametros publicados permitan reproducir
    # el pronostico- no depende de que modelo gane, y asi se comprueba.
    serie = _serie_estable(120)
    resultado = ejecutar_proyeccion(serie, 2033, 1, 2021)
    codigo = resultado["modelo_codigo"]
    parametros = resultado.get("parametros_modelo", {})
    esperados = {
        "drift": {"primer_valor", "ultimo_valor", "pendiente_mensual"},
        "lineal": {"beta_0", "beta_1"},
    }
    assert codigo in esperados, f"modelo sin parametros esperados en la prueba: {codigo}"
    assert esperados[codigo].issubset(parametros), (codigo, sorted(parametros))
    lineas = "\n".join(_lineas_parametros_reproducibles(resultado, serie))
    receta = "\n".join(_lineas_receta_reproduccion(resultado))
    # La ecuacion publicada es la del modelo entregado, con sus parametros.
    assert "y_hat" in lineas, lineas
    assert any(p in lineas for p in esperados[codigo]), lineas
    assert "CSV" in receta
    df = construir_dataframe_reproducibilidad(serie, resultado)
    assert {"periodo", "valor_observado", "valor_proyectado"}.issubset(df.columns)
    assert not [c for c in df.columns if "ic80" in str(c).lower()]
    # ENDURECIDO 17-08-2026 (V-CODEX-R3, residual 1). P0-C / C2 conservaba
    # `ic95_inferior` e `ic95_superior` vacias «porque su presencia es el contrato
    # del CSV». Se RETIRARON: ningun consumidor dependia de ellas, y una columna
    # permanentemente vacia cuyo nombre anuncia el intervalo del 95 % en la
    # cabecera es una afirmacion sobre un objeto que no se entrega. Con ellas se
    # fueron `metodo_intervalo`, `sigma_h_intervalo`, `q95_intervalo` y los dos
    # percentiles.
    #
    # La vigilancia se endurece: ahora no basta que lleguen vacias, no deben
    # existir.
    for retirada in ("ic95_inferior", "ic95_superior", "ic95_relativo",
                     "metodo_intervalo", "sigma_h_intervalo", "q95_intervalo",
                     "percentil_95_inf_intervalo", "percentil_95_sup_intervalo"):
        assert retirada not in df.columns, f"volvio al CSV la columna del intervalo: {retirada}"
    assert not [c for c in df.columns if "cobertura" in str(c).lower()], list(df.columns)
    # Y lo que SI debe seguir: los agregados fuera de muestra.
    assert {"rmse_horizonte", "mae_horizonte", "mase_horizonte"}.issubset(df.columns), list(df.columns)
    assert {
        "tipo_registro",
        "horizonte",
        "modelo_ganador_por_horizonte",
        "modelo_final_aplicado",
        "modelo_final_difiere_ganador",
        "metricas_reportadas",
        "estado_horizonte",
        "clasificacion_horizonte",
        "confianza_horizonte",
        "permitido_para_proyeccion_tecnica",
        "permitido_como_escenario",
        "no_recomendable",
        "razon_decision",
        "mensaje_no_recomendables",
        "horizonte_maximo_recomendado",
        "horizonte_maximo_con_cautela",
        "horizonte_maximo_escenario",
        "primer_horizonte_no_viable",
        # P0-C / ESTRATEGIA C2, 15-08-2026: se retira `ic95_relativo` de la lista
        # exigida. Es el ancho relativo de la banda: una magnitud del intervalo
        # que esta version ya no publica. Decirle al lector cuan ancha es una
        # banda que no recibe no le permite hacer nada con el dato y sugiere que
        # la incertidumbre si esta acotada. El calculo interno se conserva.
        "rmse_horizonte",
        "mae_horizonte",
        "mape_horizonte",
        "smape_horizonte",
        "mase_horizonte",
    }.issubset(df.columns)
    assert "decision_horizonte" in set(df["tipo_registro"])
    periodos_con_dato = df.loc[df["periodo"].astype(str).str.len() > 0, "periodo"]
    assert periodos_con_dato.str.contains("-").all()
    with tempfile.TemporaryDirectory() as tmp:
        ruta = Path(tmp) / "repro.csv"
        generar_csv_reproducibilidad(ruta, serie, resultado)
        assert ruta.exists() and ruta.stat().st_size > 0


def test_pdf_y_docx_contienen_horizontes_con_escenario() -> None:
    from app_icociv.reportes.generador_reportes import generar_reporte_pdf, generar_reporte_proyeccion
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    serie = _serie_estable(96)
    resultado = ejecutar_proyeccion(serie, 2030, 6, 2021)
    fila = pd.DataFrame({"Codigo": ["TEST"], "Descripcion": ["Serie estable"]})
    with tempfile.TemporaryDirectory() as tmp:
        ruta_pdf = Path(tmp) / "informe.pdf"
        ruta_docx = Path(tmp) / "informe.docx"
        generar_reporte_pdf(ruta_pdf, "Prueba", "fuente.xlsx", {}, {"horizonte": 18}, [], "T_16", fila, serie, resultado, [])
        generar_reporte_proyeccion(
            ruta_docx,
            "Prueba",
            "fuente.xlsx",
            {},
            {"horizonte": 18},
            [],
            "T_16",
            fila,
            serie,
            resultado,
            [],
        )
        assert ruta_pdf.exists() and ruta_pdf.stat().st_size > 0
        assert ruta_docx.exists() and ruta_docx.stat().st_size > 0
        try:
            from docx import Document

            texto = "\n".join(p.text for p in Document(str(ruta_docx)).paragraphs)
            # Títulos del informe rediseñado (27-jul-2026). El contenido
            # metodológico es el mismo; cambiaron los encabezados.
            assert "Fundamento estadístico del análisis" in texto
            assert "Horizonte estadístico admisible" in texto
            assert "Evaluación por horizonte" in texto
            assert "Criterio de selección del modelo" in texto
            assert "Parámetros y ecuación del modelo aplicado" in texto
            assert "Receta de reproducción de la proyección" in texto
            assert "Hyndman" in texto
        except ImportError:
            pass


def test_pdf_y_docx_sin_proyeccion() -> None:
    from app_icociv.reportes.generador_reportes import generar_reporte_pdf, generar_reporte_proyeccion
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    serie = pd.DataFrame({"Periodo": _periodos(10), "Indice": [100 + i for i in range(10)]})
    resultado = ejecutar_proyeccion(serie, 2022, 1, 2021)
    fila = pd.DataFrame({"Codigo": ["TEST"], "Descripcion": ["Serie corta"]})
    with tempfile.TemporaryDirectory() as tmp:
        ruta_pdf = Path(tmp) / "sin_proyeccion.pdf"
        ruta_docx = Path(tmp) / "sin_proyeccion.docx"
        generar_reporte_pdf(ruta_pdf, "Prueba", "fuente.xlsx", {}, {"horizonte": 1}, [], "T_16", fila, serie, resultado, [])
        generar_reporte_proyeccion(
            ruta_docx,
            "Prueba",
            "fuente.xlsx",
            {},
            {"horizonte": 1},
            [],
            "T_16",
            fila,
            serie,
            resultado,
            [],
        )
        assert ruta_pdf.exists() and ruta_pdf.stat().st_size > 0
        assert ruta_docx.exists() and ruta_docx.stat().st_size > 0


def test_horizonte_dinamico_con_serie_drift() -> None:
    """Caso 7777777 con la serie que selecciona Drift (fixture serie_drift.csv)."""
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion
    from app_icociv.reportes.generador_reportes import _lineas_factibilidad, _lineas_horizontes

    resultado = ejecutar_proyeccion(_serie_777_drift(), 2027, 1, 2021)
    assert resultado["horizonte_solicitado"] == 12
    assert resultado["proyeccion_generada"] is True
    assert resultado["horizonte_permitido"] == 12

    lineas_factibilidad = "\n".join(_lineas_factibilidad(resultado["factibilidad"], resultado))
    assert "No se permite proyección porque el backtesting muestra errores extremos recurrentes" not in lineas_factibilidad
    assert "Horizonte máximo permitido como escenario:" in lineas_factibilidad
    assert "Estado general: No recomendable" not in lineas_factibilidad

    # CIERRE 08-08-2026: h=18 ya no se degrada a escenario. Lo hacia V-12
    # -benchmark en h>=13-, un criterio sin fuente que no consultaba ningun
    # deterioro medido. La etiqueta del tramo se conserva; el veto desaparece.
    #
    # CIERRE H-9: el fixture se llama `serie_drift.csv` porque bajo la regla
    # ponderada por 1/h seleccionaba Drift. Con la seleccion por RMSE fuera de
    # muestra global elige otro modelo. El nombre del fixture es historico; lo
    # que la prueba fija -y era su proposito (H-13)- es que las dos series de
    # prueba recorren trayectorias distintas y no dependen de ningun archivo
    # residual, y que el informe es coherente con el modelo que se entregue.
    tabla_horizontes = "\n".join(_lineas_horizontes(resultado))
    assert "12 meses" in tabla_horizontes and "permitido=Si" in tabla_horizontes
    assert "18 meses" in tabla_horizontes and "técnico=Si" in tabla_horizontes
    assert "Escenario estadístico extendido" in tabla_horizontes
    # Un unico modelo en toda la trayectoria, y es el que el informe declara.
    visible = str(resultado["model_name"])
    filas_modelo = [linea for linea in tabla_horizontes.splitlines() if "modelo_final=" in linea]
    assert filas_modelo, tabla_horizontes
    for fila in filas_modelo:
        assert f"modelo_final={visible}" in fila, (visible, fila)


def test_horizonte_dinamico_con_serie_holt() -> None:
    """Caso 7777777 con la serie que recorre la trayectoria de Holt.

    Es el escenario que antes solo aparecia cuando faltaba el CSV residual. Se
    fija como prueba propia para que las dos trayectorias queden cubiertas
    siempre y no dependan de que un archivo exista o no (H-13).

    AUDITORIA 09-08-2026 (C-01): esta prueba exigia por su NOMBRE el modelo
    `Holt lineal`. Al pasar los coeficientes de suavizamiento de fijos a
    estimados por minimizacion del SSE (FPP3 8.1-8.2), la misma serie pasa a
    seleccionar la variante amortiguada. Su fallo fue la demostracion de que la
    parametrizacion fija cambiaba el modelo entregado.

    Se conserva lo que la prueba verificaba de fondo -que este fixture recorre
    la rama de Holt, distinta de la del fixture hermano- y se sustituye la
    identidad exacta del ganador por la propiedad. Es el mismo criterio que ya
    se aplico al fixture `serie_drift.csv` en el cierre H-9: cual sea el modelo
    depende de la regla y de la parametrizacion, y ambas pueden cambiar con
    fuente; lo que no puede cambiar es que las dos trayectorias esten cubiertas.

    CIERRE F/H/G, 13-08-2026. Quedaban dos asserts que seguian fijando el
    ganador por su NOMBRE -`model_name.startswith("Holt")` y `modelo_final=Holt`
    en la tabla-. La integracion vuelve a cambiar la evidencia disponible
    -rejilla derivada de la aritmetica de ventanas y retiro del tratamiento
    calendario- y con ella el modelo seleccionado. Fijar la identidad del
    ganador convierte un resultado empirico en contrato (REQ 31), y ademas
    contradice lo que esta misma docstring ya declaraba. Se retiran.

    La cobertura de la rama de Holt **no se pierde**: pasa a obtenerse por
    construccion en `test_la_rama_holt_se_recorre_por_construccion`, forzando el
    conjunto de candidatos, que es determinista y no depende de quien gane.
    """
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion
    from app_icociv.reportes.generador_reportes import _lineas_horizontes

    resultado = ejecutar_proyeccion(_serie_777_holt(), 2027, 1, 2021)
    assert resultado["horizonte_solicitado"] == 12
    assert resultado["proyeccion_generada"] is True
    assert resultado["horizonte_permitido"] == 12

    tabla_horizontes = chr(10).join(_lineas_horizontes(resultado))
    assert "12 meses" in tabla_horizontes and "permitido=Si" in tabla_horizontes
    # El informe declara el modelo REALMENTE entregado, sea cual sea. Es el mismo
    # criterio del fixture hermano: coherencia informe-resultado, no identidad.
    visible = str(resultado["model_name"])
    filas_modelo = [linea for linea in tabla_horizontes.splitlines() if "modelo_final=" in linea]
    assert filas_modelo, tabla_horizontes
    for fila in filas_modelo:
        assert f"modelo_final={visible}" in fila, (visible, fila)


def test_la_rama_holt_se_recorre_por_construccion() -> None:
    """Cobertura de la rama de Holt sin depender de que Holt gane la seleccion.

    `ajustar_modelos_candidatos` y `ejecutar_backtesting_comparativo` aceptan el
    conjunto de candidatos. Restringiendolo a una unica variante de Holt, la rama
    se recorre **matematicamente**: el resultado no depende de la regla de
    seleccion, de la parametrizacion ni de que serie se use. Esto es lo que la
    prueba hermana pretendia cubrir cuando exigia el nombre del ganador.
    """
    from app_icociv.estadistica.modelos_interpretables import (
        MODELOS_SERIE_TEMPORAL,
        ajustar_modelos_candidatos,
        proyectar_modelo,
    )
    from app_icociv.validacion.backtesting import ejecutar_backtesting_comparativo

    serie = _serie_777_holt()
    y = serie["Indice"].to_numpy(dtype=float)
    t = np.arange(len(y), dtype=float)

    for nombre in ("holt_lineal", "holt_amortiguado"):
        assert nombre in MODELOS_SERIE_TEMPORAL
        candidatos = ajustar_modelos_candidatos(t, y, modelos=(nombre,))
        assert len(candidatos) == 1, candidatos
        ajuste = candidatos[0]
        assert not ajuste.get("error"), (nombre, ajuste.get("error"))
        # La rama produce una trayectoria finita para los pasos pedidos.
        futuro = proyectar_modelo(ajuste, np.arange(len(y), len(y) + 12, dtype=float))
        assert len(futuro) == 12
        assert np.all(np.isfinite(futuro)), (nombre, futuro)

    # Y la rama tambien se recorre dentro del backtesting con candidato unico,
    # que es donde se mide la evidencia fuera de muestra.
    banco = ejecutar_backtesting_comparativo(
        serie[["Periodo", "Indice"]], modelos=("holt_amortiguado",), horizontes=(1, 3), anio_base=2021
    )
    # El banco se indexa por "{modelo}_h{h}": el candidato unico debe ser el
    # unico modelo presente, en los dos horizontes pedidos.
    assert set(banco) == {"holt_amortiguado_h1", "holt_amortiguado_h3"}, sorted(banco)
    assert all((banco[clave] or {}).get("ejecutado") for clave in banco), banco


def test_el_limite_60_es_operativo_de_entrada_y_no_techo_estadistico() -> None:
    """P0-H: 60 acota lo que el usuario puede PEDIR, no lo que la evidencia permite.

    La reclasificacion se verifica por dos vias independientes, ninguna basada en
    leer un comentario:

    1. **alcance**: el unico consumidor de la constante es la validacion del dato
       de entrada, de modo que rechaza 61 y acepta 60;
    2. **no gobierna la rejilla**: el techo evaluado sale de la aritmetica de
       ventanas, y con una serie suficientemente larga la rejilla supera 60 sin
       que la constante intervenga.
    """
    from app_icociv.estadistica.criterios import MIN_ITERACIONES_WF_ESCENARIO
    from app_icociv.proyeccion.servicio_proyeccion import (
        HORIZONTE_MAXIMO_OPERATIVO,
        _limites_auditoria_horizontes,
        validar_horizonte_solicitado,
    )

    # 1. Limite de ENTRADA.
    assert validar_horizonte_solicitado(HORIZONTE_MAXIMO_OPERATIVO) == HORIZONTE_MAXIMO_OPERATIVO
    try:
        validar_horizonte_solicitado(HORIZONTE_MAXIMO_OPERATIVO + 1)
    except Exception:
        pass
    else:
        raise AssertionError("el limite operativo debe rechazar la entrada fuera de rango")

    # 2. No gobierna la rejilla: el techo evaluable lo fija la aritmetica de
    #    ventanas, y con serie larga supera 60 sin que la constante intervenga.
    #    Se comprueba sobre la funcion que FIJA el techo, no proyectando 100+
    #    horizontes: es la misma propiedad y no cuesta veinte minutos.
    entrenamiento, maximo_por_datos, limite, _ = _limites_auditoria_horizontes(120)
    assert maximo_por_datos > HORIZONTE_MAXIMO_OPERATIVO, maximo_por_datos
    assert limite == maximo_por_datos, (limite, maximo_por_datos)
    # P0-G, 16-08-2026 (V-CODEX-3). Aqui se exigia que el techo dejara
    # MIN_ITERACIONES_WF_ESCENARIO = 3 ventanas. Ese tres procede -por
    # declaracion de su propia ficha- de estimar la dispersion y verificar la
    # cobertura del INTERVALO, eje que P0-C retiro del producto: un requisito de
    # la BANDA estaba recortando los horizontes en que se entrega el PUNTO.
    #
    # El techo pasa a ser la cota de EXISTENCIA, `h <= n - N0`, que no es un
    # umbral elegido sino la frontera del dato: con cero ventanas no hay ningun
    # error fuera de muestra. Se comprueba POR AMBOS LADOS, de modo que la
    # prueba fija la igualdad exacta y no una desigualdad que cualquier techo
    # mas conservador satisfaria.
    ventanas = lambda h: 120 - entrenamiento - h + 1  # noqa: E731
    assert ventanas(maximo_por_datos) >= 1, ventanas(maximo_por_datos)
    assert ventanas(maximo_por_datos + 1) < 1, ventanas(maximo_por_datos + 1)
    # Y la constante sobrevive como corte DESCRIPTIVO: existe, pero no recorta.
    assert MIN_ITERACIONES_WF_ESCENARIO >= 1
    assert maximo_por_datos > 120 - entrenamiento - MIN_ITERACIONES_WF_ESCENARIO + 1, (
        "el techo ya no puede depender del piso de ventanas"
    )


def test_las_dos_series_del_caso_777_son_distintas_y_dan_trayectorias_distintas() -> None:
    """Deja constancia de por que H-13 era un problema y no una casualidad.

    Lo que importa es que **los dos fixtures son series distintas y producen
    resultados distintos**, de modo que la suite cubre las dos trayectorias sin
    depender de que un CSV residual exista en el disco.

    CIERRE H-9, 08-08-2026: cual sea cada modelo depende de la regla de
    seleccion, y la regla cambio. Los nombres de los fixtures son historicos.

    AUDITORIA 09-08-2026 (C-01): esta prueba exigia ademas que los dos modelos
    fueran DISTINTOS ENTRE SI. Al estimar los coeficientes de Holt en lugar de
    fijarlos, ambos fixtures pasan a seleccionar la misma variante amortiguada,
    y la prueba fallo. Es una consecuencia real y se deja declarada: con sus
    parametros estimados, Holt gana en las dos series sinteticas.

    La identidad del ganador NO era el proposito de la prueba. Lo que H-13
    exigia es que las dos ramas se recorran de verdad y que el resultado no
    dependa de un archivo residual. Eso se fija ahora sobre las TRAYECTORIAS,
    que es la propiedad que importa y que no puede coincidir si las series
    difieren.
    """
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion

    drift, holt = _serie_777_drift(), _serie_777_holt()
    assert len(drift) == len(holt)
    assert not drift["Indice"].equals(holt["Indice"]), "Los fixtures deben diferir"

    uno = ejecutar_proyeccion(drift, 2027, 1, 2021)
    dos = ejecutar_proyeccion(holt, 2027, 1, 2021)
    assert uno["model_name"] and dos["model_name"]
    assert uno["proyeccion_generada"] is True and dos["proyeccion_generada"] is True

    valores_uno = list(uno["proyecciones"]["indice_proyectado"])
    valores_dos = list(dos["proyecciones"]["indice_proyectado"])
    assert valores_uno and valores_dos
    assert valores_uno != valores_dos, (
        "Los dos fixtures deben producir trayectorias distintas; si coinciden, "
        "la suite estaria cubriendo una sola rama (H-13)."
    )


def test_ninguna_prueba_depende_de_reportes_generados() -> None:
    """Hermeticidad: este archivo no puede leer carpetas de salida."""
    fuente = Path(__file__).read_text(encoding="utf-8-sig")
    cuerpo = fuente.split("def test_ninguna_prueba_depende_de_reportes_generados")[0]
    assert 'ROOT / "reportes_generados"' not in cuerpo
    assert "reportes_generados" not in cuerpo.replace(
        "reportes_generados/7777777.csv **si", ""
    ), "Las pruebas no deben leer de reportes_generados"


def test_intervalos_con_sesgo_son_simetricos_y_mas_anchos() -> None:
    """Con errores de un solo signo la banda se desplaza pero contiene el pronostico."""
    from app_icociv.proyeccion.servicio_proyeccion import _intervalos_prediccion

    predicciones = np.array([100.0, 101.0])
    errores_por_horizonte = {
        1: np.array([-1.0, -2.0, -3.0, -4.0, -5.0]),
        2: np.array([-1.5, -2.5, -3.5, -4.5, -5.5]),
    }
    intervalos = _intervalos_prediccion(
        y_futuro=predicciones,
        errores_por_horizonte=errores_por_horizonte,
    )
    primero = intervalos[0]
    assert primero["limite_inferior_95"] < predicciones[0]
    assert primero["limite_superior_95"] > predicciones[0]
    # AUDITORIA 09-08-2026 (P0-C). Antes se exigia que la banda se abriera hacia
    # el lado del sesgo, porque el intervalo se centraba en la media del error.
    # FPP3 5.4 situa el remedio del sesgo en el PRONOSTICO -«simply add m to all
    # forecasts»-, no en el intervalo. La banda es ahora simetrica alrededor del
    # pronostico; el sesgo se absorbe ENSANCHANDOLA, porque sigma_h usa SUM e^2.
    izquierda = predicciones[0] - primero["limite_inferior_95"]
    derecha = primero["limite_superior_95"] - predicciones[0]
    assert abs(izquierda - derecha) < 1e-9, (izquierda, derecha)
    assert int(primero["errores_oos_disponibles"]) == 5


def test_mase_mayor_uno_no_redacta_inferior_si_es_auxiliar() -> None:
    from app_icociv.estadistica.analisis_series import MENSAJE_MASE_MAYOR_UNO
    from app_icociv.validacion.backtesting import interpretar_backtesting

    texto = MENSAJE_MASE_MAYOR_UNO + " " + interpretar_backtesting(
        {"mape": 1.2, "mase": 1.4, "rmse": 1.0, "mae": 0.8, "estabilidad_error": 0.4}
    )
    assert "inferior al benchmark naive" not in texto
    assert "criterio de descarte" in texto


def test_docx_777_separa_advertencias_globales_y_horizonte() -> None:
    from docx import Document
    from app_icociv.proyeccion.servicio_proyeccion import ejecutar_proyeccion
    from app_icociv.reportes.generador_reportes import generar_reporte_proyeccion

    serie = _serie_777_drift()
    resultado = ejecutar_proyeccion(serie, 2027, 1, 2021)
    fila = pd.DataFrame({"Caso": ["7777777"], "Descripcion": ["Caso con h=12 permitido y h=18 no recomendable"]})
    with tempfile.TemporaryDirectory() as tmp:
        ruta_docx = Path(tmp) / "7777777_corregido.docx"
        generar_reporte_proyeccion(
            ruta_docx,
            "Prueba",
            "7777777.csv",
            {},
            {"horizonte": 12},
            [],
            "T_16",
            fila,
            serie,
            resultado,
            [],
        )
        texto = "\n".join(
            [p.text for p in Document(str(ruta_docx)).paragraphs]
            + ["\t".join(cell.text for cell in row.cells) for table in Document(str(ruta_docx)).tables for row in table.rows]
        )
    assert "No se permite proyección porque el backtesting muestra errores extremos recurrentes" not in texto
    # El informe rediseñado separa la clasificacion (columna de la tabla de
    # horizontes) del motivo (viñeta por clasificacion). Ambos deben estar.
    filas_h18 = [linea for linea in texto.splitlines() if linea.startswith("18" + chr(9))]
    assert filas_h18, "Falta la fila h=18 en la evaluacion por horizonte."
    # CIERRE 08-08-2026: h=18 conserva una etiqueta que nombra el TRAMO de
    # horizonte y pierde la de «alta incertidumbre», que era el veto de V-12 y
    # de los cortes de amplitud. La fila sigue publicandose con sus metricas.
    assert "alta incertidumbre" not in filas_h18[0].lower(), filas_h18[0]
    # P0-G, 13-08-2026. Antes se exigia que la clasificacion de h=18 fuera
    # «Escenario estadistico extendido» o «Proyeccion extendida»: otro conjunto
    # cerrado de peldanos. Retirada la escalera, la clasificacion es descriptiva.
    # Lo sustantivo -y lo que el comentario de arriba ya declaraba- es que la fila
    # se siga PUBLICANDO con sus metricas y sin la etiqueta del veto retirado.
    celdas = filas_h18[0].split(chr(9))
    assert len(celdas) >= 5, filas_h18[0]
    assert celdas[-1].strip(), filas_h18[0]
    numeros = [c for c in celdas[1:-1] if any(ch.isdigit() for ch in c)]
    assert len(numeros) >= 3, filas_h18[0]
    for retirada in ("alta incertidumbre", "no recomendable", "no se permite"):
        assert retirada not in filas_h18[0].lower(), (retirada, filas_h18[0])
    # P0-G: la frase «El horizonte h=12 se permite como ...» solo se emitia en la
    # rama de `_estado_por_horizonte` correspondiente a dos peldanos de la
    # escalera. Retirada esta, la redaccion es otra. Lo que la prueba promete por
    # su nombre -y lo que se conserva- es que el informe SEPARE lo global de lo
    # que es propio de cada horizonte, y que declare el permiso del solicitado.
    filas_h12 = [linea for linea in texto.splitlines() if linea.startswith("12" + chr(9))]
    assert filas_h12, "Falta la fila h=12 en la evaluacion por horizonte."
    assert filas_h12[0] != filas_h18[0]
    assert int(resultado["horizonte_permitido"]) == 12, resultado["horizonte_permitido"]
    assert "se permite" in texto.lower(), texto[:400]
    # CIERRE F/H/G: se retira «Drift» como ganador exigido por su nombre. El
    # fixture se llama `serie_drift.csv` por razones historicas y el propio
    # archivo ya declaro en H-9 que el ganador depende de la regla y de la
    # parametrizacion. Lo que el informe debe cumplir es coherencia: declarar el
    # modelo que de verdad se entrego.
    assert "Modelo aplicado" in texto
    assert str(resultado["model_name"]) in texto, resultado["model_name"]


def main() -> None:
    test_horizonte_18_se_permite_como_escenario()
    test_horizonte_25_estable_se_permite_si_hay_evidencia()
    test_horizonte_25_erratico_no_se_fuerza()
    test_umbral_h12_31_es_alta_incertidumbre_no_bloqueo()
    test_umbral_h18_535_no_recomienda()
    test_horizonte_maximo_no_es_18_si_h18_no_cumple()
    test_horizontes_se_evaluan_mes_a_mes_y_no_solo_atajos_ui()
    test_mensaje_no_recomendables_no_imprime_lista_vacia()
    test_h18_no_se_redacta_como_intervalos_aceptables()
    test_el_modelo_entregado_es_el_de_minimo_error_oos_global()
    test_combo_no_cambia_con_rueda_y_periodo_iso()
    test_grafica_usa_periodos_reales()
    test_contenedor_ruta_tiene_scroll_y_tooltip_hover_combo()
    test_el_catalogo_se_define_por_estimabilidad_y_no_por_longitud()
    test_no_se_aplican_ensambles_de_modelos()
    test_mase_auxiliar_y_primer_periodo_proyectado()
    test_graficas_reporte_y_fundamento_estadistico()
    test_modelos_tradicionales_y_no_ejecutados_quedan_trazados()
    test_parametros_drift_receta_y_csv_reproducible()
    test_pdf_y_docx_contienen_horizontes_con_escenario()
    test_pdf_y_docx_sin_proyeccion()
    test_horizonte_dinamico_con_serie_drift()
    test_horizonte_dinamico_con_serie_holt()
    test_la_rama_holt_se_recorre_por_construccion()
    test_el_limite_60_es_operativo_de_entrada_y_no_techo_estadistico()
    test_las_dos_series_del_caso_777_son_distintas_y_dan_trayectorias_distintas()
    test_ninguna_prueba_depende_de_reportes_generados()
    test_intervalos_con_sesgo_son_simetricos_y_mas_anchos()
    test_mase_mayor_uno_no_redacta_inferior_si_es_auxiliar()
    test_docx_777_separa_advertencias_globales_y_horizonte()
    print("Pruebas de horizonte dinámico, UI y reportes OK")


if __name__ == "__main__":
    main()
