"""Renderizado de un :class:`Informe` a documento Word editable.

El DOCX es el formato de trabajo: el usuario debe poder cambiar textos, ajustar
tablas, añadir conclusiones y firmar. Por eso todo se construye con estilos
nombrados de Word y tablas reales, nunca con imágenes de texto ni cuadros
posicionados a mano.
"""

from __future__ import annotations

import io
from typing import Any

from app_icociv.reportes.modelo import (
    Aviso,
    Ficha,
    Firmas,
    Formula,
    FUENTE_MONO,
    FUENTE_TEXTO,
    Imagen,
    Informe,
    NOMBRE_COMPLETO,
    PALETA,
    Parrafo,
    Portada,
    Seccion,
    Tabla,
    Vinetas,
    fecha_larga,
)

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    DOCX_DISPONIBLE = True
except ImportError:  # pragma: no cover - depende del entorno
    Document = None  # type: ignore
    DOCX_DISPONIBLE = False


# Umbral de secciones a partir del cual el documento supera las ocho páginas y
# merece índice automático (§8.1). El informe ejecutivo queda por debajo.
SECCIONES_PARA_INDICE = 10
ANCHO_UTIL_CM = 16.0


def esta_disponible() -> bool:
    return DOCX_DISPONIBLE


def _color(hex_color: str) -> Any:
    return RGBColor.from_string(hex_color.lstrip("#").upper())


def _sombrear(celda: Any, hex_color: str) -> None:
    sombreado = OxmlElement("w:shd")
    sombreado.set(qn("w:val"), "clear")
    sombreado.set(qn("w:fill"), hex_color.lstrip("#").upper())
    celda._tc.get_or_add_tcPr().append(sombreado)


def _borde_parrafo(parrafo: Any, hex_color: str, lado: str = "left", grosor: int = 18) -> None:
    pPr = parrafo._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    borde = OxmlElement(f"w:{lado}")
    borde.set(qn("w:val"), "single")
    borde.set(qn("w:sz"), str(grosor))
    borde.set(qn("w:space"), "8")
    borde.set(qn("w:color"), hex_color.lstrip("#").upper())
    bordes.append(borde)
    pPr.append(bordes)


def _fondo_parrafo(parrafo: Any, hex_color: str) -> None:
    pPr = parrafo._p.get_or_add_pPr()
    sombreado = OxmlElement("w:shd")
    sombreado.set(qn("w:val"), "clear")
    sombreado.set(qn("w:fill"), hex_color.lstrip("#").upper())
    pPr.append(sombreado)


def _campo(parrafo: Any, instruccion: str, texto_provisional: str = "") -> None:
    """Inserta un campo de Word (PAGE, NUMPAGES, TOC) que Word recalcula al abrir."""
    corrida = parrafo.add_run()
    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruccion
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    provisional = OxmlElement("w:t")
    provisional.text = texto_provisional
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    for elemento in (inicio, instr, separador, provisional, fin):
        corrida._r.append(elemento)


def _repetir_encabezado(fila: Any) -> None:
    trPr = fila._tr.get_or_add_trPr()
    encabezado = OxmlElement("w:tblHeader")
    encabezado.set(qn("w:val"), "true")
    trPr.append(encabezado)


def _no_partir_fila(fila: Any) -> None:
    trPr = fila._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _crear_estilos(documento: Any) -> None:
    """Define los estilos SAVIP. Al ser estilos con nombre, el usuario los edita
    una vez y el cambio se propaga a todo el documento."""
    estilos = documento.styles

    normal = estilos["Normal"]
    normal.font.name = FUENTE_TEXTO
    normal.font.size = Pt(10)
    normal.font.color.rgb = _color(PALETA["texto"])
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    jerarquia = (("Title", 24, True), ("Heading 1", 16, True), ("Heading 2", 13, True), ("Heading 3", 11.5, True))
    for nombre, tamano, negrita in jerarquia:
        try:
            estilo = estilos[nombre]
        except KeyError:  # pragma: no cover - plantilla atípica
            continue
        estilo.font.name = FUENTE_TEXTO
        estilo.font.size = Pt(tamano)
        estilo.font.bold = negrita
        estilo.font.color.rgb = _color(PALETA["marca_intensa"] if nombre != "Title" else PALETA["marca"])
        estilo.paragraph_format.space_before = Pt(14 if nombre == "Heading 1" else 10)
        estilo.paragraph_format.space_after = Pt(6)
        estilo.paragraph_format.keep_with_next = True

    from docx.enum.style import WD_STYLE_TYPE

    def _nuevo(nombre: str, base: str = "Normal") -> Any:
        try:
            return estilos[nombre]
        except KeyError:
            estilo = estilos.add_style(nombre, WD_STYLE_TYPE.PARAGRAPH)
            estilo.base_style = estilos[base]
            return estilo

    pie_figura = _nuevo("SAVIP Pie de figura")
    pie_figura.font.size = Pt(8.5)
    pie_figura.font.italic = True
    pie_figura.font.color.rgb = _color(PALETA["texto_secundario"])
    pie_figura.paragraph_format.space_before = Pt(2)

    nota = _nuevo("SAVIP Nota")
    nota.font.size = Pt(8.5)
    nota.font.color.rgb = _color(PALETA["texto_secundario"])

    formula = _nuevo("SAVIP Formula")
    formula.font.name = FUENTE_MONO
    formula.font.size = Pt(9.5)
    formula.paragraph_format.space_after = Pt(2)
    formula.paragraph_format.left_indent = Cm(0.8)

    aviso = _nuevo("SAVIP Aviso")
    aviso.font.size = Pt(9.5)
    aviso.paragraph_format.left_indent = Cm(0.4)
    aviso.paragraph_format.space_after = Pt(3)


def _configurar_pagina(documento: Any, informe: Informe) -> None:
    for seccion in documento.sections:
        seccion.top_margin = Cm(2.2)
        seccion.bottom_margin = Cm(2.0)
        seccion.left_margin = Cm(2.5)
        seccion.right_margin = Cm(2.5)
        _encabezado_pie(seccion, informe)


def _encabezado_pie(seccion: Any, informe: Informe) -> None:
    encabezado = seccion.header.paragraphs[0]
    encabezado.text = ""
    encabezado.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    corrida = encabezado.add_run(f"{NOMBRE_COMPLETO}  ·  {informe.identificador}")
    corrida.font.size = Pt(8)
    corrida.font.color.rgb = _color(PALETA["texto_secundario"])

    pie = seccion.footer.paragraphs[0]
    pie.text = ""
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefijo = pie.add_run(f"{fecha_larga(informe.generado)}  ·  Página ")
    prefijo.font.size = Pt(8)
    prefijo.font.color.rgb = _color(PALETA["texto_secundario"])
    _campo(pie, " PAGE ", "1")
    intermedio = pie.add_run(" de ")
    intermedio.font.size = Pt(8)
    intermedio.font.color.rgb = _color(PALETA["texto_secundario"])
    _campo(pie, " NUMPAGES ", "1")
    for corrida in pie.runs:
        corrida.font.size = Pt(8)
        corrida.font.color.rgb = _color(PALETA["texto_secundario"])


def _escribir_portada(documento: Any, portada: Portada, informe: Informe) -> None:
    if portada.logo:
        parrafo = documento.add_paragraph()
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            parrafo.add_run().add_picture(io.BytesIO(portada.logo), width=Cm(4.5))
        except Exception:  # pragma: no cover - logo inválido no debe romper el informe
            pass

    documento.add_paragraph().add_run()  # respiro superior
    titulo = documento.add_paragraph(portada.titulo, style="Title")
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitulo = documento.add_paragraph()
    corrida = subtitulo.add_run(portada.subtitulo)
    corrida.font.size = Pt(12)
    corrida.font.color.rgb = _color(PALETA["texto_secundario"])
    _borde_parrafo(subtitulo, PALETA["marca"], "left", 24)

    documento.add_paragraph()
    _tabla_clave_valor(documento, portada.filas, ancho_clave=5.4)

    if portada.observaciones:
        documento.add_paragraph("Observaciones generales", style="Heading 3")
        documento.add_paragraph(portada.observaciones)

    if informe.secciones_visibles():
        documento.add_page_break()


def _escribir_indice(documento: Any) -> None:
    documento.add_paragraph("Contenido", style="Heading 1")
    parrafo = documento.add_paragraph()
    # El campo TOC lo resuelve Word al abrir o al pulsar F9; así el índice sigue
    # siendo correcto después de que el usuario edite el documento.
    _campo(parrafo, r' TOC \o "1-3" \h \z \u ', "Actualice el campo (F9) para generar el índice.")
    documento.add_page_break()


def _tabla_clave_valor(documento: Any, filas: list[tuple[str, str]], ancho_clave: float = 5.0) -> Any:
    tabla = documento.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = False
    for indice, (campo, valor) in enumerate(filas):
        fila = tabla.add_row()
        _no_partir_fila(fila)
        celda_campo, celda_valor = fila.cells
        celda_campo.width = Cm(ancho_clave)
        celda_valor.width = Cm(ANCHO_UTIL_CM - ancho_clave)
        _escribir_celda(celda_campo, str(campo), negrita=True)
        _escribir_celda(celda_valor, str(valor))
        if indice % 2 == 0:
            _sombrear(celda_campo, PALETA["superficie_alterna"])
            _sombrear(celda_valor, PALETA["superficie_alterna"])
        else:
            _sombrear(celda_campo, PALETA["superficie"])
    return tabla


def _escribir_celda(celda: Any, texto: str, negrita: bool = False, derecha: bool = False, tamano: float = 9.5) -> None:
    celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    parrafo = celda.paragraphs[0]
    parrafo.text = ""
    parrafo.paragraph_format.space_after = Pt(2)
    parrafo.paragraph_format.space_before = Pt(2)
    if derecha:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    corrida = parrafo.add_run(texto)
    corrida.font.size = Pt(tamano)
    corrida.font.bold = negrita


def _escribir_tabla(documento: Any, bloque: Tabla) -> None:
    if bloque.titulo:
        parrafo = documento.add_paragraph()
        corrida = parrafo.add_run(bloque.titulo)
        corrida.font.bold = True
        corrida.font.size = Pt(9.5)
        corrida.font.color.rgb = _color(PALETA["marca_intensa"])
        parrafo.paragraph_format.keep_with_next = True

    columnas = len(bloque.encabezados)
    tabla = documento.add_table(rows=1, cols=columnas)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = False

    anchos = bloque.anchos or tuple([ANCHO_UTIL_CM / columnas] * columnas)
    escala = ANCHO_UTIL_CM / sum(anchos) if sum(anchos) else 1.0
    anchos = tuple(a * escala for a in anchos)

    encabezado = tabla.rows[0]
    _repetir_encabezado(encabezado)
    _no_partir_fila(encabezado)
    for indice, titulo in enumerate(bloque.encabezados):
        celda = encabezado.cells[indice]
        celda.width = Cm(anchos[indice])
        _escribir_celda(celda, titulo, negrita=True, derecha=indice in bloque.columnas_numericas, tamano=9.0)
        _sombrear(celda, PALETA["marca"])
        for parrafo in celda.paragraphs:
            for corrida in parrafo.runs:
                corrida.font.color.rgb = _color("#FFFFFF")

    for numero, valores in enumerate(bloque.filas):
        fila = tabla.add_row()
        _no_partir_fila(fila)
        for indice in range(columnas):
            celda = fila.cells[indice]
            celda.width = Cm(anchos[indice])
            texto = str(valores[indice]) if indice < len(valores) else ""
            _escribir_celda(celda, texto, derecha=indice in bloque.columnas_numericas, tamano=9.0)
            if numero % 2 == 1:
                _sombrear(celda, PALETA["superficie_alterna"])

    for texto, etiqueta in ((bloque.nota, "Nota"), (bloque.fuente, "Fuente")):
        if texto:
            documento.add_paragraph(f"{etiqueta}: {texto}", style="SAVIP Nota")


def _escribir_ficha(documento: Any, bloque: Ficha) -> None:
    if bloque.destacados:
        tabla = documento.add_table(rows=2, cols=len(bloque.destacados))
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.autofit = False
        ancho = ANCHO_UTIL_CM / len(bloque.destacados)
        for indice, (etiqueta, valor) in enumerate(bloque.destacados):
            celda_etiqueta = tabla.rows[0].cells[indice]
            celda_valor = tabla.rows[1].cells[indice]
            celda_etiqueta.width = Cm(ancho)
            celda_valor.width = Cm(ancho)
            _escribir_celda(celda_etiqueta, etiqueta.upper(), tamano=7.5)
            _escribir_celda(celda_valor, valor, negrita=True, tamano=12.0)
            _sombrear(celda_etiqueta, PALETA["marca_suave"])
            _sombrear(celda_valor, PALETA["marca_suave"])
            for parrafo in celda_valor.paragraphs:
                for corrida in parrafo.runs:
                    corrida.font.color.rgb = _color(PALETA["marca_intensa"])
        _no_partir_fila(tabla.rows[0])
        _no_partir_fila(tabla.rows[1])
        documento.add_paragraph()
    if bloque.filas:
        _tabla_clave_valor(documento, bloque.filas, ancho_clave=6.0)


def _escribir_aviso(documento: Any, bloque: Aviso) -> None:
    colores = {
        "advertencia": (PALETA["aviso"], PALETA["aviso_fondo"], "!"),
        "error": (PALETA["error"], PALETA["error_fondo"], "×"),
        "informacion": (PALETA["informacion"], PALETA["informacion_fondo"], "i"),
    }
    borde, fondo, simbolo = colores.get(bloque.nivel, colores["advertencia"])

    titulo = documento.add_paragraph(style="SAVIP Aviso")
    corrida = titulo.add_run(f"{simbolo}  {bloque.titulo}")
    corrida.font.bold = True
    corrida.font.color.rgb = _color(borde)
    _borde_parrafo(titulo, borde, "left", 24)
    _fondo_parrafo(titulo, fondo)
    titulo.paragraph_format.keep_with_next = True

    for item in bloque.items:
        parrafo = documento.add_paragraph(style="SAVIP Aviso")
        parrafo.add_run(f"—  {item}")
        _borde_parrafo(parrafo, borde, "left", 24)
        _fondo_parrafo(parrafo, fondo)


def _escribir_formula(documento: Any, bloque: Formula) -> None:
    etiqueta = documento.add_paragraph()
    corrida = etiqueta.add_run(bloque.etiqueta)
    corrida.font.bold = True
    corrida.font.size = Pt(9.5)
    etiqueta.paragraph_format.keep_with_next = True
    documento.add_paragraph(bloque.general, style="SAVIP Formula")
    for linea in bloque.sustitucion:
        documento.add_paragraph(linea, style="SAVIP Formula")
    resultado = documento.add_paragraph(bloque.resultado, style="SAVIP Formula")
    for corrida in resultado.runs:
        corrida.font.bold = True
        corrida.font.color.rgb = _color(PALETA["marca_intensa"])


def _escribir_firmas(documento: Any, bloque: Firmas) -> None:
    documento.add_paragraph()
    tabla = documento.add_table(rows=2, cols=len(bloque.roles))
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    ancho = ANCHO_UTIL_CM / len(bloque.roles)
    for indice, rol in enumerate(bloque.roles):
        firma = tabla.rows[0].cells[indice]
        etiqueta = tabla.rows[1].cells[indice]
        firma.width = Cm(ancho)
        etiqueta.width = Cm(ancho)
        _escribir_celda(firma, "\n\n")
        _escribir_celda(etiqueta, f"{rol}\nNombre, cargo y fecha", tamano=8.5)
        borde = OxmlElement("w:tcBorders")
        superior = OxmlElement("w:top")
        superior.set(qn("w:val"), "single")
        superior.set(qn("w:sz"), "8")
        superior.set(qn("w:color"), PALETA["borde_fuerte"].lstrip("#").upper())
        borde.append(superior)
        etiqueta._tc.get_or_add_tcPr().append(borde)


def _escribir_bloque(documento: Any, bloque: Any) -> None:
    if isinstance(bloque, Parrafo):
        parrafo = documento.add_paragraph(bloque.texto)
        if bloque.enfasis:
            for corrida in parrafo.runs:
                corrida.font.bold = True
    elif isinstance(bloque, Vinetas):
        for item in bloque.items:
            documento.add_paragraph(item, style="List Bullet")
    elif isinstance(bloque, Ficha):
        _escribir_ficha(documento, bloque)
    elif isinstance(bloque, Tabla):
        _escribir_tabla(documento, bloque)
    elif isinstance(bloque, Imagen):
        parrafo = documento.add_paragraph()
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parrafo.add_run().add_picture(io.BytesIO(bloque.datos), width=Cm(min(bloque.ancho_cm, ANCHO_UTIL_CM)))
        if bloque.pie:
            pie = documento.add_paragraph(bloque.pie, style="SAVIP Pie de figura")
            pie.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif isinstance(bloque, Aviso):
        _escribir_aviso(documento, bloque)
    elif isinstance(bloque, Formula):
        _escribir_formula(documento, bloque)
    elif isinstance(bloque, Firmas):
        _escribir_firmas(documento, bloque)


def _escribir_seccion(documento: Any, seccion: Seccion, numero: int) -> None:
    documento.add_paragraph(f"{numero}. {seccion.titulo}", style="Heading 1")
    for bloque in seccion.bloques:
        _escribir_bloque(documento, bloque)


def construir_documento(informe: Informe) -> Any:
    """Devuelve el objeto ``Document`` ya compuesto."""
    if not DOCX_DISPONIBLE:
        raise RuntimeError("python-docx no está instalado. Ejecute: pip install python-docx")

    documento = Document()
    _crear_estilos(documento)
    _configurar_pagina(documento, informe)

    documento.core_properties.title = informe.identificador
    documento.core_properties.subject = NOMBRE_COMPLETO
    documento.core_properties.comments = f"Informe {informe.tipo} generado por SAVIP."

    if informe.portada is not None:
        _escribir_portada(documento, informe.portada, informe)

    secciones = informe.secciones_visibles()
    if len(secciones) >= SECCIONES_PARA_INDICE:
        _escribir_indice(documento)

    for numero, seccion in enumerate(secciones, start=1):
        _escribir_seccion(documento, seccion, numero)

    return documento


def guardar(informe: Informe, ruta: Any) -> Any:
    from pathlib import Path

    destino = Path(ruta)
    if destino.suffix.lower() != ".docx":
        destino = destino.with_suffix(".docx")
    destino.parent.mkdir(parents=True, exist_ok=True)
    construir_documento(informe).save(str(destino))
    return destino


def a_bytes(informe: Informe) -> bytes:
    memoria = io.BytesIO()
    construir_documento(informe).save(memoria)
    return memoria.getvalue()
